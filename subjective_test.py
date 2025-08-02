# -*- coding: utf-8 -*-
import sys
import json
from docx import Document
from docx2pdf import convert


def main(input_json):
    # Load data
    with open(input_json, 'r', encoding='utf-8') as f:
        data = json.load(f)

    heading = data['heading']
    questions = data['questions']
    output_filename = data.get('output_filename', 'Questions')

    # Generate DOCX
    doc = Document()
    doc.add_heading(heading, level=1)

    for i, q in enumerate(questions, 1):
        doc.add_paragraph(f"{i}. {q}")

    docx_file = f"{output_filename}.docx"
    pdf_file = f"{output_filename}.pdf"
    doc.save(docx_file)

    # Convert to PDF
    convert(docx_file, pdf_file)
    print(f"Generated: {docx_file} and {pdf_file}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python generate_questions.py resources/filename.json")
    else:
        main(sys.argv[1])
