# -*- coding: utf-8 -*-
import json
import sys
from docx import Document

# Read filename from command line or use default
if len(sys.argv) < 2:
    print("❌ Please provide a JSON file from the 'resources/' directory.")
    print("Usage: python3 matsci_mcq.py resources/prog_fundamentals.json")
    sys.exit(1)

json_path = sys.argv[1]

# Load JSON data
with open(json_path, "r", encoding="utf-8") as f:
    data = json.load(f)

title = data.get("title", "MCQ Test")
file_name = data.get("file_name", "MCQ_Document.docx")
mcq_data = data.get("mcq_data", [])

# Create document
doc = Document()
doc.add_heading(title, level=1)

# Generate MCQs
for idx, mcq in enumerate(mcq_data, 1):
    options = mcq.get("options", [])
    correct_index_1based = mcq.get("correct", 1)  # default to first option if missing

    table = doc.add_table(rows=4 + 4, cols=3)  # 4 options + 4 metadata rows
    table.style = 'Table Grid'

    # Question
    table.cell(0, 0).text = f"Question #{idx}"
    table.cell(0, 1).merge(table.cell(0, 2)).text = mcq.get("question", "")

    # Question type
    table.cell(1, 0).text = "Type"
    table.cell(1, 1).merge(table.cell(1, 2)).text = "multiple_choice"

    # Options
    for i, opt in enumerate(options):
        table.cell(2 + i, 0).text = "Option"
        table.cell(2 + i, 1).text = opt
        # compare with 1-based index from JSON
        table.cell(2 + i, 2).text = "correct" if (i + 1) == correct_index_1based else "incorrect"

    # Solution
    table.cell(6, 0).text = "Solution"
    table.cell(6, 1).merge(table.cell(6, 2)).text = mcq.get("solution", "")

    # Marks
    table.cell(7, 0).text = "Marks"
    table.cell(7, 1).text = "1"
    table.cell(7, 2).text = "0"

    doc.add_paragraph("")

# Save file
doc.save(file_name)
print(f"✅ Document created successfully: {file_name}")
